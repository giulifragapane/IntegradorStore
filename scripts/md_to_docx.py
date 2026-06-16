"""Convert GUION-VIDEO-STORE-APP.md to DOCX."""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text: str, base_bold=False, base_italic=False):
    """Parse **bold** and `code` inline markers."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.bold = base_bold
            run.italic = base_italic
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
            run.italic = base_italic
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.bold = base_bold
            run.italic = base_italic
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold
        run.italic = base_italic


def add_paragraph(doc, text: str, style=None, bold=False, italic=False, indent=0):
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    add_formatted_runs(p, text, base_bold=bold, base_italic=italic)
    return p


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def convert_md_to_docx(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=0)
            h.runs[0].font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            i += 1
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:])
                i += 1
            text = " ".join(quote_lines)
            p = add_paragraph(doc, text, italic=True, indent=0.25)
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            headers = parse_table_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for j, h in enumerate(headers):
                hdr_cells[j].text = h
                set_cell_shading(hdr_cells[j], "D9E2F3")
                for p in hdr_cells[j].paragraphs:
                    for run in p.runs:
                        run.bold = True
            for r_idx, row in enumerate(rows):
                row_cells = table.rows[r_idx + 1].cells
                for j, cell in enumerate(row):
                    if j < len(row_cells):
                        row_cells[j].text = cell
            doc.add_paragraph()
            continue

        if re.match(r"^\d+\.\s", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                item = re.sub(r"^\d+\.\s", "", lines[i].strip())
                add_paragraph(doc, item, style="List Number")
                i += 1
            continue

        if stripped.startswith("- "):
            while i < len(lines) and lines[i].strip().startswith("- "):
                item = lines[i].strip()[2:]
                add_paragraph(doc, item, style="List Bullet")
                i += 1
            continue

        if stripped.startswith("**[") and stripped.endswith("]**"):
            p = add_paragraph(doc, stripped.strip("*"), bold=True)
            p.runs[0].font.color.rgb = RGBColor(0x70, 0x50, 0x00)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and ":" in stripped:
            add_paragraph(doc, stripped.strip("*"), bold=True)
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.save(docx_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    md_file = root / "GUION-VIDEO-STORE-APP.md"
    out_file = root / "GUION-VIDEO-STORE-APP.docx"
    convert_md_to_docx(md_file, out_file)
    print(f"Created: {out_file}")
